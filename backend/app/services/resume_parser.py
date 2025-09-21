import json
from typing import Dict, List, Any, Optional
import PyPDF2
import docx
import re
from datetime import datetime

class ResumeParser:
    @staticmethod
    def parse_resume(file_path: str) -> Dict[str, Any]:
        try:
            text = ResumeParser.extract_text(file_path)
            if not text.strip():
                return {
                    'content': '',
                    'skills': [],
                    'experience': [],
                    'education': []
                }
            
            # Clean and normalize text
            cleaned_text = ResumeParser.clean_text(text)
            
            # Parse all sections
            skills = ResumeParser.extract_skills(cleaned_text)
            experience = ResumeParser.extract_experience(cleaned_text)
            education = ResumeParser.extract_education(cleaned_text)
            
            return {
                'content': cleaned_text,
                'skills': skills,
                'experience': experience,
                'education': education
            }
            
        except Exception as e:
            print(f"Error parsing resume: {e}")
            return {
                'content': '',
                'skills': [],
                'experience': [],
                'education': []
            }

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize resume text"""
        # Remove extra whitespace and normalize line breaks
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        text = text.strip()
        return text

    @staticmethod
    def extract_skills(text: str) -> List[str]:
        """Extract skills from resume text with multiple strategies"""
        skills = set()
        text_lower = text.lower()
        
        # Strategy 1: Common technical skills
        technical_skills = [
            'python', 'javascript', 'java', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
            'html', 'css', 'sass', 'less', 'react', 'angular', 'vue', 'svelte',
            'node.js', 'express', 'django', 'flask', 'fastapi', 'spring', 'laravel',
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible',
            'git', 'jenkins', 'circleci', 'github actions', 'jira', 'confluence',
            'machine learning', 'deep learning', 'ai', 'tensorflow', 'pytorch',
            'data analysis', 'pandas', 'numpy', 'tableau', 'power bi',
            'project management', 'agile', 'scrum', 'kanban'
        ]
        
        for skill in technical_skills:
            if skill in text_lower:
                skills.add(skill.title())
        
        # Strategy 2: Look for skills section with various patterns
        skills_patterns = [
            r'(?i)(?:technical\s+)?skills(?:\s+&?\s+competencies)?[:\n]*(.*?)(?=\n\n|\n(?:experience|education|projects|$))',
            r'(?i)competencies[:\n]*(.*?)(?=\n\n|\n(?:experience|education|projects|$))',
            r'(?i)technical\s+expertise[:\n]*(.*?)(?=\n\n|\n(?:experience|education|projects|$))'
        ]
        
        for pattern in skills_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                skills_text = match.group(1)
                # Extract skills from bullet points, commas, or new lines
                bullet_points = re.findall(r'[•\-*]\s*([^\n•\-*]+)', skills_text)
                comma_separated = re.findall(r'[^,\n]+(?:\s*,\s*[^,\n]+)*', skills_text)
                
                for item in bullet_points + comma_separated:
                    item = item.strip()
                    if item and len(item) > 2 and not item.lower().startswith(('years', 'level')):
                        skills.add(item.strip())
        
        # Strategy 3: Look for technologies mentioned throughout the resume
        tech_keywords = re.findall(r'\b(?:[A-Z][a-z]*\+*[A-Z]*[a-z]*\d*|React\.js|Node\.js|Angular\.js)\b', text)
        for tech in tech_keywords:
            if len(tech) > 2 and tech.lower() not in ['the', 'and', 'for', 'with']:
                skills.add(tech)
        
        return sorted(list(skills))

    @staticmethod
    def extract_experience(text: str) -> List[Dict[str, str]]:
        """Extract work experience with better pattern matching"""
        experience = []
        
        # Multiple patterns to find experience section
        experience_patterns = [
            r'(?i)(?:work\s+)?experience(?:\s+history)?[:\n]*(.*?)(?=\n\n|\n(?:education|skills|projects|$))',
            r'(?i)employment\s+history[:\n]*(.*?)(?=\n\n|\n(?:education|skills|projects|$))',
            r'(?i)professional\s+experience[:\n]*(.*?)(?=\n\n|\n(?:education|skills|projects|$))'
        ]
        
        for pattern in experience_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                exp_text = match.group(1)
                break
        else:
            # If no experience section found, return empty
            return experience
        
        # Split into individual job entries
        job_entries = re.split(r'\n(?=\s*(?:[A-Z][a-z]+\s+\d{4}\s*[-–]|Present|Current|[\w\s]+\s+\d{4}))', exp_text)
        
        for entry in job_entries:
            if not entry.strip():
                continue
                
            # Extract job details
            lines = [line.strip() for line in entry.split('\n') if line.strip()]
            
            if len(lines) >= 2:
                # Try to extract title and company from first line
                first_line = lines[0]
                title_company_match = re.match(r'(.+?)\s*(?:at|@|,)\s*(.+)', first_line)
                
                if title_company_match:
                    title = title_company_match.group(1).strip()
                    company = title_company_match.group(2).strip()
                else:
                    title = first_line
                    company = lines[1] if len(lines) > 1 else ''
                
                # Extract duration (look for date patterns)
                duration = ''
                for line in lines:
                    date_match = re.search(r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\s*[-–]\s*(?:Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})|\d{4}\s*[-–]\s*(?:Present|Current|\d{4}))', line, re.IGNORECASE)
                    if date_match:
                        duration = date_match.group(1)
                        break
                
                # Extract description (remaining lines)
                description_lines = []
                for line in lines[1:]:  # Skip first line (title/company)
                    if not re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|Present|Current|\d{4}\s*[-–]', line):
                        description_lines.append(line)
                
                description = ' '.join(description_lines)
                
                experience.append({
                    'title': title,
                    'company': company,
                    'duration': duration,
                    'description': description
                })
        
        return experience

    @staticmethod
    def extract_education(text: str) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        
        # Multiple patterns to find education section
        education_patterns = [
            r'(?i)education[:\n]*(.*?)(?=\n\n|\n(?:experience|skills|projects|$))',
            r'(?i)academic\s+background[:\n]*(.*?)(?=\n\n|\n(?:experience|skills|projects|$))',
            r'(?i)qualifications[:\n]*(.*?)(?=\n\n|\n(?:experience|skills|projects|$))'
        ]
        
        for pattern in education_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                edu_text = match.group(1)
                break
        else:
            return education
        
        # Split into individual education entries
        edu_entries = re.split(r'\n(?=\s*(?:[A-Z]|\d{4}))', edu_text)
        
        for entry in edu_entries:
            if not entry.strip():
                continue
                
            lines = [line.strip() for line in entry.split('\n') if line.strip()]
            
            if lines:
                # Try to extract degree and institution
                first_line = lines[0]
                degree_institution_match = re.match(r'(.+?)\s*(?:,|at|@)\s*(.+)', first_line)
                
                if degree_institution_match:
                    degree = degree_institution_match.group(1).strip()
                    institution = degree_institution_match.group(2).strip()
                else:
                    degree = first_line
                    institution = lines[1] if len(lines) > 1 else ''
                
                # Extract year/date
                year = ''
                for line in lines:
                    year_match = re.search(r'\b(?:19|20)\d{2}\b', line)
                    if year_match:
                        year = year_match.group()
                        break
                
                education.append({
                    'degree': degree,
                    'institution': institution,
                    'year': year
                })
        
        return education

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from PDF or DOCX files"""
        if file_path.endswith('.pdf'):
            return ResumeParser.extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return ResumeParser.extract_text_from_docx(file_path)
        else:
            return ""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF with better handling"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
        return text

    @staticmethod
    def debug_parsing(file_path: str):
        """Debug function to see what's being parsed"""
        text = ResumeParser.extract_text(file_path)
        print("=== EXTRACTED TEXT ===")
        print(text[:1000] + "..." if len(text) > 1000 else text)
        print("\n=== SKILLS ===")
        skills = ResumeParser.extract_skills(text)
        print(skills)
        print("\n=== EXPERIENCE ===")
        experience = ResumeParser.extract_experience(text)
        print(experience)
        print("\n=== EDUCATION ===")
        education = ResumeParser.extract_education(text)
        print(education)